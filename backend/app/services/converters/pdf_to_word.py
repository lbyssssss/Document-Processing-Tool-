# PDF to Word Conversion
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from typing import Optional
from pathlib import Path

class PDFToWordConverter:
    """PDF转Word转换器"""

    def __init__(self, pdf_path: Path, output_path: Path):
        self.pdf_path = pdf_path
        self.output_path = output_path

    def convert(self) -> dict:
        """将PDF转换为Word格式"""
        try:
            # 读取PDF文件
            reader = PdfReader(str(self.pdf_path))

            # 创建Word文档
            doc = Document()
            doc.add_heading('PDF转换文档', 0)

            # 提取每页的文本内容
            page_content = ""
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    # 添加页面标题
                    doc.add_heading(f'第{i + 1}页', level=1)

                    # 添加页面内容
                    # 将文本分割成段落
                    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                    for para in paragraphs:
                        doc.add_paragraph(para)

                    # 添加分页符（除了最后一页）
                    if i < len(reader.pages) - 1:
                        doc.add_page_break()

            # 保存Word文档
            doc.save(str(self.output_path))

            return {
                "success": True,
                "page_count": len(reader.pages),
                "output_file": str(self.output_path),
                "message": "Word文档生成成功",
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
            }

    def extract_metadata(self) -> dict:
        """提取PDF元数据"""
        try:
            reader = PdfReader(str(self.pdf_path))
            metadata = reader.metadata

            return {
                "title": metadata.get("/Title", "") if metadata else "",
                "author": metadata.get("/Author", "") if metadata else "",
                "subject": metadata.get("/Subject", "") if metadata else "",
                "creator": metadata.get("/Creator", "") if metadata else "",
                "creation_date": metadata.get("/CreationDate", "") if metadata else "",
            }

        except Exception as e:
            return {
                "error": str(e),
            }
