# Word to PDF Conversion
from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from typing import Optional

class WordToPDFConverter:
    """Word转PDF转换器"""

    def __init__(self, word_path: Path, output_path: Path):
        self.word_path = word_path
        self.output_path = output_path

    def extract_content(self) -> dict:
        """从Word文档中提取内容"""
        try:
            doc = Document(str(self.word_path))

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格数据
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text if cell.text else "")
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                "success": True,
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
            }

    def convert(self) -> dict:
        """将Word转换为PDF格式"""
        try:
            # 首先提取Word内容
            content = self.extract_content()
            if not content.get("success"):
                return content

            # 创建PDF文档
            pdf = SimpleDocTemplate(str(self.output_path), pagesize=letter)

            # 添加Word段落到PDF
            styles = self._get_sample_styles()

            story = []

            for i, para_text in enumerate(content.get("paragraphs", [])):
                p = Paragraph(para_text, styles["BodyText"])
                story.append(p)
                # 添加段落间空格
                story.append(Spacer(1, 0.2 * Inch))

            # 添加表格
            for table in content.get("tables", []):
                self._add_table_to_pdf(story, table)

            # 一次性build所有内容
            pdf.build(story)

            return {
                "success": True,
                "output_file": str(self.output_path),
                "paragraphs_converted": len(content.get("paragraphs", [])),
                "tables_converted": len(content.get("tables", [])),
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
            }

    def _get_sample_styles(self) -> dict:
        """获取示例样式"""
        from reportlab.lib.styles import getSampleStyleSheet
        styles = getSampleStyleSheet()

        return {
            "Normal": styles["Normal"],
            "Title": styles["Title"],
            "Heading1": styles["Heading1"],
            "BodyText": styles["BodyText"],
        }

    def _add_table_to_pdf(self, story: list, table_data: list) -> None:
        """添加表格到PDF"""
        from reportlab.platypus import Table, TableStyle

        if not table_data or not table_data[0]:
            return

        # 创建表格样式
        style = TableStyle([
            ('BACKGROUND', (0, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), colors.black),
            ('ALIGN', (0, 0), 'LEFT'),
            ('FONTNAME', (0, 0), 'Helvetica'),
            ('FONTSIZE', (0, 0), 10),
            ('BOTTOMPADDING', (0, 0), 12),
            ('BACKGROUND', (1, 0), colors.beige),
            ('GRID', (0, 1), 1, colors.black),
        ])

        # 创建表格
        t = Table(table_data)
        t.setStyle(style)
        t.wrapOn(pdf, 6.5 * Inch, 9.5 * Inch)
        story.append(t)
